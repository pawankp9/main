from docx import Document
import os

def read_word_file(file_path):
    
    # Replace 'your_file.docx' with the actual path to your Word file
    file_name_without_extension, _ = os.path.splitext(os.path.basename(file_path))
    document = Document(file_path)

    fileContent = ""
    # Iterate through paragraphs in the document
    for paragraph in document.paragraphs:
        # print(f"1Content: {paragraph.text}")
        fileContent += paragraph.text

    # Iterate through tables in the document (if any)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                # print(f"2Content: {cell.text}")
                fileContent += cell.text
    
    return fileContent, file_name_without_extension


